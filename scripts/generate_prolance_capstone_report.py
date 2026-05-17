# -*- coding: utf-8 -*-
"""
Build Prolance capstone-style Word report from AfetHat template styles.
- Copies the example .docx, clears body only (keeps styles, theme, footer).
- Re-inserts template logo image1.png without editing the asset.
- Adds diagrams (PNG) and tables; run: python scripts/generate_prolance_capstone_report.py
"""
from __future__ import annotations

import io
import shutil
import sys
import zipfile
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402
from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Inches, Pt  # noqa: E402

REPO = Path(__file__).resolve().parents[1]
OUT_DOCX = REPO / "docs" / "Prolance_Capstone_Report.docx"
FIG_DIR = REPO / "docs" / "report_generated"
TEMPLATE_CANDIDATES = [
    Path(r"c:\Users\ozgur\Downloads\AfetHat_Capstone_Report_son.docx"),
    REPO.parent / "AfetHat_Capstone_Report_son.docx",
]

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def find_template() -> Path | None:
    for p in TEMPLATE_CANDIDATES:
        if p.is_file():
            return p
    return None


def extract_logo_png(template: Path) -> Path | None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    out = FIG_DIR / "template_logo_image1.png"
    with zipfile.ZipFile(template, "r") as zf:
        name = "word/media/image1.png"
        if name not in zf.namelist():
            return None
        out.write_bytes(zf.read(name))
    return out


def make_figure_context(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(7, 4.2))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")
    boxes = [
        (1, 3.5, "Client\n(Flutter / Web)"),
        (4, 3.5, "Freelancer\n(Flutter / Web)"),
        (7, 3.5, "Admin\n(Next.js)"),
        (2.5, 1.2, "Supabase\nPostgres · Auth · Storage · RPC"),
        (6.5, 1.2, "Next.js\nMarketing · Portal"),
    ]
    for x, y, t in boxes:
        ax.add_patch(
            plt.Rectangle(
                (x - 0.85, y - 0.45),
                1.7,
                0.9,
                fill=True,
                facecolor="#E8F0FE",
                edgecolor="#174EA6",
                linewidth=1.2,
            )
        )
        ax.text(x, y, t, ha="center", va="center", fontsize=9, wrap=True)
    ax.annotate(
        "",
        xy=(4, 2.95),
        xytext=(4, 2.1),
        arrowprops=dict(arrowstyle="-|>", color="#333"),
    )
    ax.text(5.0, 2.45, "JWT / RLS", fontsize=8, color="#333")
    ax.set_title("Figure 1 — System context (high level)", fontsize=11, fontweight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=160)
    plt.close(fig)


def make_figure_architecture(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(7.2, 4.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis("off")
    layers = [
        (5, 6.2, "Presentation\nFlutter client · Web build", "#FFF3E0", "#E65100"),
        (5, 4.5, "Application services\nGoRouter · Provider · Repositories", "#E8F5E9", "#2E7D32"),
        (5, 2.8, "Backend\nSupabase Postgres + RPC + Storage (deliverables)", "#E3F2FD", "#1565C0"),
        (5, 1.1, "Web portal\nNext.js App Router (portal / admin)", "#F3E5F5", "#6A1B9A"),
    ]
    for x, y, text, fc, ec in layers:
        ax.add_patch(
            plt.Rectangle((1.2, y - 0.55), 7.6, 1.05, facecolor=fc, edgecolor=ec, linewidth=1.2)
        )
        ax.text(x, y - 0.02, text, ha="center", va="center", fontsize=9.5)
    for y in [5.65, 3.95, 2.25]:
        ax.annotate(
            "",
            xy=(5, y - 0.55),
            xytext=(5, y + 0.5),
            arrowprops=dict(arrowstyle="-|>", color="#444", lw=1),
        )
    ax.set_title("Figure 2 — Layered architecture", fontsize=11, fontweight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=160)
    plt.close(fig)


def make_figure_escrow_flow(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(7.2, 2.8))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 2)
    ax.axis("off")
    steps = [
        (1.2, "Proposal\nsubmitted"),
        (3.4, "Client accepts\nescrow HELD"),
        (5.6, "Delivery files\n+ lifecycle"),
        (7.8, "Client review\nawaiting_client_review"),
        (10.0, "Accept →\npayout_pending"),
    ]
    xs = [s[0] for s in steps]
    for i, (x, label) in enumerate(steps):
        ax.add_patch(
            plt.Circle((x, 1), 0.55, facecolor="#E3F2FD", edgecolor="#0D47A1", linewidth=1.1)
        )
        ax.text(x, 1, label, ha="center", va="center", fontsize=7.5)
    for a, b in zip(xs, xs[1:]):
        ax.annotate("", xy=(b - 0.55, 1), xytext=(a + 0.55, 1), arrowprops=dict(arrowstyle="-|>", lw=1.2))
    ax.set_title("Figure 3 — Contract / escrow lifecycle (simplified)", fontsize=11, fontweight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=160)
    plt.close(fig)


def clear_body_keep_sectpr(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag == f"{W_NS}sectPr":
            continue
        body.remove(child)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str | None = None):
    if caption:
        cap = doc.add_paragraph(caption, style="Body Text")
        cap.runs[0].italic = True
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Normal Table"
    except KeyError:
        table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            table.cell(r_i, c_i).text = val


def add_figure(doc: Document, png: Path, caption: str, width_in: float = 5.8):
    doc.add_picture(str(png), width=Inches(width_in))
    p = doc.add_paragraph(caption, style="Body Text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.italic = True
        r.font.size = Pt(10)


def main() -> int:
    template = find_template()
    if template is None:
        print("Template AfetHat_Capstone_Report_son.docx not found.", file=sys.stderr)
        return 1

    FIG_DIR.mkdir(parents=True, exist_ok=True)
    fig_ctx = FIG_DIR / "fig1_context.png"
    fig_arch = FIG_DIR / "fig2_layers.png"
    fig_esc = FIG_DIR / "fig3_escrow.png"
    make_figure_context(fig_ctx)
    make_figure_architecture(fig_arch)
    make_figure_escrow_flow(fig_esc)

    logo = extract_logo_png(template)
    shutil.copy(template, OUT_DOCX)
    doc = Document(str(OUT_DOCX))
    clear_body_keep_sectpr(doc)

    # --- Title block (structure mirrors example; text length kept in similar bands) ---
    if logo:
        doc.add_picture(str(logo), width=Inches(2.15))
        doc.add_paragraph()

    t = doc.add_paragraph(
        "OSTIM TECHNICAL UNIVERSITY FACULTY OF ENGINEERING CAPSTONE PROJECT REPORT",
        style="Body Text",
    )
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in t.runs:
        r.bold = True

    doc.add_paragraph()
    title = doc.add_paragraph(
        "Prolance: A Cross-Platform Freelance Marketplace with Escrow-Backed Contracts, "
        "Supabase-Backed Data Plane, and Dual Web Surfaces (Flutter · Android / iOS / Web · Next.js)",
        style="Title",
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        "Design and Implementation of a Demo-Ready Monorepo Linking a Flutter Client, "
        "a Next.js Marketing and Employer Portal, and Managed PostgreSQL with Row-Level Security",
        style="Body Text",
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.italic = True

    for _ in range(3):
        doc.add_paragraph()

    doc.add_paragraph("Project Team", style="Body Text").alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in (
        "[Student Name Surname] ([Student Number])",
        "[Student Name Surname] ([Student Number])",
        "[Student Name Surname] ([Student Number])",
    ):
        p = doc.add_paragraph(line, style="Body Text")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph("Project Advisor", style="Heading 4").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Prof. Dr. [Advisor Name Surname]", style="Body Text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(4):
        doc.add_paragraph()

    foot = doc.add_paragraph(
        "Undergraduate Project — Department of Computer Engineering — MFBP 402(1) Graduation Project II",
        style="Body Text",
    )
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("February 2026", style="Body Text").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- Front matter (same order as example) ---
    doc.add_paragraph("APPROVAL", style="Heading 1")
    doc.add_paragraph()
    doc.add_paragraph(
        "This undergraduate capstone project, entitled “Prolance: A Cross-Platform Freelance Marketplace with "
        "Escrow-Backed Contracts, Supabase-Backed Data Plane, and Dual Web Surfaces”, prepared by the project team "
        "listed on the title page, has been examined and is accepted as a capstone project report in partial fulfilment "
        "of the requirements for the Bachelor of Science degree in Computer Engineering.",
        style="Body Text",
    )
    doc.add_paragraph()
    for label in ("Supervisor", "Member", "Member", "Head of Department"):
        doc.add_paragraph(label, style="Heading 4")
        doc.add_paragraph("[Name Surname]", style="Body Text")
        doc.add_paragraph()

    doc.add_page_break()
    doc.add_paragraph("ACKNOWLEDGEMENTS", style="Heading 1")
    doc.add_paragraph()
    for para in (
        "This report documents the analysis, design, implementation, and verification of the Prolance platform: "
        "a monorepository that couples a Flutter client for job discovery, proposals, messaging, and escrow-aware "
        "contract progression with a Supabase project (PostgreSQL, Auth, Storage, RPC) and a Next.js surface for "
        "marketing, authentication, and an employer-facing portal including dispute review for administrators.",
        "The implementation extends an initial UI-centric prototype into a working vertical slice: authenticated "
        "profiles, job postings, proposal lifecycles, private deliverable storage, signed download links, and "
        "server-side transitions that keep financial state consistent with row-level security policies.",
        "We thank the project advisor for structured feedback on architecture and risk management, and peers who "
        "participated in exploratory testing on web builds and mobile emulators.",
    ):
        doc.add_paragraph(para, style="Body Text")

    doc.add_paragraph()
    doc.add_paragraph("Prolance Team — February 2026", style="Body Text")
    doc.add_paragraph("Ankara", style="Body Text")

    doc.add_page_break()
    doc.add_paragraph("ÖZET", style="Heading 1")
    doc.add_paragraph(
        "Serbest çalışan pazar yerlerinde güven eksikliği, ödeme anlaşmazlıkları ve dağınık iletişim kanalları "
        "işveren ile freelancer arasında sürtünme yaratır. Bu çalışma, Prolance adlı çapraz platform uygulamasını "
        "tasarlayarak bu sürtünmeyi azaltmayı hedefler: Flutter istemcisi iş ilanlarını keşfetme, teklif verme, "
        "mesajlaşma ve teslimat dosyalarını güvenli depoda toplama akışlarını sunar; Supabase tarafında Postgres "
        "üzerinde satır düzeyi güvenlik, RPC ile yaşam döngüsü geçişleri ve özel depolama ile teslim dosyalarının "
        "yalnızca taraflarca okunması sağlanır. Next.js ile pazarlama ve portal yüzeyleri birleştirilir; yönetici "
        "paneli anlaşmazlık incelemeleri için ek ekranlar içerir. Çalışma, Agile sprint düzeni, risk kaydı ve "
        "otomatikleştirilmiş testlerle desteklenen bir mühendislik sürecini belgeler.",
        style="Body Text",
    )
    doc.add_paragraph(
        "Anahtar kelimeler: serbest çalışma platformu, Flutter, Supabase, satır düzeyi güvenlik, emanet ödeme, "
        "Next.js, portal, teslimat yaşam döngüsü.",
        style="Body Text",
    )

    doc.add_page_break()
    doc.add_paragraph("ABSTRACT", style="Heading 1")
    doc.add_paragraph()
    doc.add_paragraph(
        "Freelance marketplaces must combine discoverability, negotiation, and post-award execution with trust "
        "mechanisms that survive partial failures and malicious clients. Prolance implements a pragmatic slice of "
        "this problem space: employers publish jobs, freelancers submit proposals, employers accept proposals to "
        "create escrow-backed records, freelancers upload deliverables to a private bucket, and employers review "
        "downloads before releasing funds into a timed payout window with an explicit dispute path for administrators.",
        style="Body Text",
    )
    doc.add_paragraph(
        "Keywords: freelance marketplace, Flutter, Supabase, row-level security, escrow, Next.js, employer portal, "
        "deliverable lifecycle.",
        style="Body Text",
    )

    doc.add_page_break()
    doc.add_paragraph("TABLE OF CONTENTS", style="Heading 1")
    doc.add_paragraph()
    toc_lines = [
        "Introduction ......................................................... 3",
        "Problem Definition ................................................... 5",
        "Project Planning and Management ...................................... 7",
        "System Analysis ...................................................... 9",
        "System Design ......................................................... 12",
        "Implementation ...................................................... 15",
        "Testing ............................................................. 18",
        "Working Product and MVP .............................................. 20",
        "Conclusion and Future Work ........................................... 22",
        "References ........................................................... 23",
        "Appendices ........................................................... 24",
    ]
    for line in toc_lines:
        doc.add_paragraph(line, style="List Paragraph")

    doc.add_page_break()
    doc.add_paragraph("LIST OF FIGURES", style="Heading 1")
    doc.add_paragraph()
    for cap in (
        "Figure 1 — System context (high level)",
        "Figure 2 — Layered architecture",
        "Figure 3 — Contract / escrow lifecycle (simplified)",
    ):
        doc.add_paragraph(cap, style="Body Text")

    doc.add_page_break()
    doc.add_paragraph("LIST OF TABLES", style="Heading 1")
    doc.add_paragraph()
    for cap in (
        "Table 2.1 — Limitations of informal hiring channels (comparative)",
        "Table 2.2 — Target personas (client, freelancer, admin)",
        "Table 4.1 — Functional requirements (summary)",
        "Table 4.2 — Non-functional requirements (summary)",
        "Table 5.1 — Core relational tables (summary)",
        "Table 8.1 — Working product capability matrix",
    ):
        doc.add_paragraph(cap, style="Body Text")

    doc.add_page_break()
    doc.add_paragraph("ABBREVIATIONS", style="Heading 1")
    doc.add_paragraph()
    add_table(
        doc,
        ["Abbreviation", "Meaning"],
        [
            ("RLS", "Row-Level Security (PostgreSQL)"),
            ("RPC", "Remote procedure call exposed via Supabase"),
            ("JWT", "JSON Web Token used by Supabase Auth"),
            ("MVP", "Minimum viable product / demo-ready slice"),
            ("FR / NFR", "Functional / non-functional requirement"),
        ],
    )

    # --- Introduction ---
    doc.add_page_break()
    doc.add_paragraph("Introduction", style="Heading 2")
    doc.add_paragraph("Purpose and Scope", style="Heading 3")
    doc.add_paragraph(
        "The purpose of this document is to present the engineering rationale, architecture, and verification evidence "
        "for Prolance, following the reporting conventions and section ordering of the faculty capstone template. "
        "Scope is intentionally bounded to the shipped monorepo: `client/` Flutter application, `web/` Next.js "
        "application, and `supabase/` schema, policies, and migrations that power the hosted demo project.",
        style="Body Text",
    )
    doc.add_paragraph("Motivation", style="Heading 3")
    doc.add_paragraph(
        "Informal freelance hiring fragments negotiation across chat applications, e-mail, and generic file sharing. "
        "Payment disputes are difficult to adjudicate without an auditable state machine and consistent access control "
        "to deliverables. Prolance concentrates post-award execution inside a database-backed lifecycle with explicit "
        "phases and server-side transitions, reducing ambiguity for both parties.",
        style="Body Text",
    )
    doc.add_paragraph("Structure of the Report", style="Heading 3")
    doc.add_paragraph(
        "Section 2 defines the problem and personas. Section 3 summarises planning, tooling, and risks. Section 4 "
        "captures analysis artefacts (requirements and context). Section 5 presents design views including data "
        "modelling. Section 6 details implementation surfaces. Section 7 outlines testing. Section 8 demonstrates the "
        "working product path. Section 9 concludes with lessons and future work.",
        style="Body Text",
    )

    # --- Problem definition ---
    doc.add_page_break()
    doc.add_paragraph("Problem Definition", style="Heading 2")
    doc.add_paragraph("Sectoral Background", style="Heading 3")
    doc.add_paragraph(
        "Digital talent marketplaces span general-purpose gig platforms and vertical specialist networks. Employers "
        "seek predictable delivery checkpoints; freelancers seek payment assurance and dispute visibility. Regulatory "
        "and tax treatment varies by jurisdiction; this project adopts a demonstrative escrow model suitable for "
        "academic evaluation rather than production financial compliance.",
        style="Body Text",
    )
    doc.add_paragraph("Deficiencies of Existing Systems", style="Heading 3")
    add_table(
        doc,
        ["Channel / pattern", "Typical limitation"],
        [
            ("Chat-only coordination", "No authoritative lifecycle; ambiguous “done” state."),
            ("Ad-hoc file links", "No least-privilege read model; revocable access is manual."),
            ("Informal milestone payments", "Disagreements lack structured dispute metadata."),
        ],
        caption="Table 2.1 — Limitations of informal hiring channels (comparative)",
    )
    doc.add_paragraph()
    doc.add_paragraph("Target User Profiles", style="Heading 3")
    add_table(
        doc,
        ["Persona", "Goals", "Pain points"],
        [
            (
                "Employer (client)",
                "Publish jobs, compare bids, fund escrow, review files",
                "Trust, clarity on acceptance criteria, payout timing",
            ),
            (
                "Freelancer",
                "Discover work, submit proposals, upload deliverables",
                "Payment assurance, scope creep, fragmented tooling",
            ),
            (
                "Administrator",
                "Resolve disputes, audit escalations",
                "Needs structured case notes and reproducible actions",
            ),
        ],
        caption="Table 2.2 — Target personas (client, freelancer, admin)",
    )
    doc.add_paragraph()
    doc.add_paragraph("Project Goals and Success Criteria", style="Heading 3")
    doc.add_paragraph(
        "Success is defined as a reproducible demo: two authenticated roles can progress a single job from proposal "
        "acceptance through deliverable review and timed payout simulation, with storage access enforced by policy "
        "rather than client-side checks alone.",
        style="Body Text",
    )

    # --- Planning ---
    doc.add_page_break()
    doc.add_paragraph("Project Planning and Management", style="Heading 2")
    doc.add_paragraph("Method: Agile and Sprint Planning", style="Heading 3")
    doc.add_paragraph(
        "Work was organised in short iterations with a prioritised backlog: schema and RLS foundations first, "
        "followed by Flutter repository integration, then portal routes for employers, and finally admin dispute UX.",
        style="Body Text",
    )
    doc.add_paragraph("Team Structure and Roles", style="Heading 3")
    add_table(
        doc,
        ["Role", "Responsibilities (illustrative)"],
        [
            ("Mobile / Flutter", "Routing, state, repositories, PDF and media UX"),
            ("Backend / Supabase", "Migrations, RPC, policies, storage buckets"),
            ("Web / Next.js", "Portal contracts, admin disputes, server actions"),
        ],
        caption="Table 3.1 — Team members and areas of responsibility (fill names)",
    )
    doc.add_paragraph("Version Control and Workflow", style="Heading 3")
    doc.add_paragraph(
        "The repository is hosted on GitHub (`OzgurBuyukikiz01/Prolance_Freelance_App`). Feature branches and pull "
        "requests keep migrations reviewable; Supabase changes are applied with versioned SQL files under "
        "`supabase/migrations/`.",
        style="Body Text",
    )
    doc.add_paragraph("Risk Management", style="Heading 3")
    add_table(
        doc,
        ["Risk", "Mitigation"],
        [
            ("RLS misconfiguration leaks rows", "Policy tests in SQL; least privilege by default"),
            ("Large binary assets in Git", "Deliverables only in Supabase Storage"),
            ("Web / mobile skew", "Shared Supabase contract as single source of truth"),
        ],
        caption="Table 3.3 — Risk register (abridged)",
    )

    # --- Analysis ---
    doc.add_page_break()
    doc.add_paragraph("System Analysis", style="Heading 2")
    doc.add_paragraph("Actors and System Boundaries", style="Heading 3")
    add_figure(doc, fig_ctx, "Figure 1 — System context (high level)")
    doc.add_paragraph("Functional Requirements", style="Heading 3")
    add_table(
        doc,
        ["ID", "Requirement (summary)"],
        [
            ("FR-1", "Users authenticate via Supabase Auth to access protected views."),
            ("FR-2", "Employers create and list jobs; freelancers browse and apply."),
            ("FR-3", "Employers accept proposals; system creates escrow transaction rows."),
            ("FR-4", "Freelancers register deliverables; employers download via signed URLs."),
            ("FR-5", "Employers accept or decline delivery; lifecycle advances server-side."),
            ("FR-6", "Administrators review disputes with auditable notes (portal)."),
        ],
        caption="Table 4.1 — Functional requirements (summary)",
    )
    doc.add_paragraph("Non-Functional Requirements", style="Heading 3")
    add_table(
        doc,
        ["ID", "Requirement (summary)"],
        [
            ("NFR-1", "Security: RLS on proposals, jobs, messages, and storage policies."),
            ("NFR-2", "Maintainability: typed Next.js server actions; Dart repositories."),
            ("NFR-3", "Operability: Makefile targets for local dev and deploy docs."),
        ],
        caption="Table 4.2 — Non-functional requirements (summary)",
    )
    doc.add_paragraph("Context Diagram", style="Heading 3")
    doc.add_paragraph(
        "Figure 1 situates the Flutter client and Next.js surfaces against Supabase-managed persistence. External "
        "actors are end users; the database enforces participation constraints on reads and writes.",
        style="Body Text",
    )

    # --- Design ---
    doc.add_page_break()
    doc.add_paragraph("System Design", style="Heading 2")
    doc.add_paragraph("High-Level Architecture", style="Heading 3")
    add_figure(doc, fig_arch, "Figure 2 — Layered architecture")
    doc.add_paragraph("Database Design", style="Heading 3")
    add_table(
        doc,
        ["Table / area", "Role"],
        [
            ("profiles", "User attributes, balances, admin flag"),
            ("jobs", "Listings with client ownership"),
            ("proposals", "Bid, status, lifecycle_phase, escrow-related fields"),
            ("proposal_deliveries", "File metadata and storage paths"),
            ("escrow_transactions", "Held / released / disputed states"),
            ("messages / conversations", "Threaded chat between participants"),
        ],
        caption="Table 5.1 — Core relational tables (summary)",
    )
    doc.add_paragraph("Authentication and Authorization", style="Heading 3")
    doc.add_paragraph(
        "Authentication is delegated to Supabase. Authorisation combines Postgres RLS policies with RPC entry points "
        "that validate actor roles (freelancer vs job owner) before mutating lifecycle columns.",
        style="Body Text",
    )
    doc.add_paragraph("API Design", style="Heading 3")
    doc.add_paragraph(
        "The mobile client primarily uses Supabase client libraries for CRUD and `rpc(...)` for guarded transitions "
        "(for example client delivery review and freelancer submission confirmation). The portal uses Next.js server "
        "actions with the service role only where strictly necessary, documented in repository guidelines.",
        style="Body Text",
    )
    doc.add_paragraph("User Interface Design", style="Heading 3")
    doc.add_paragraph(
        "Flutter uses a shared theme with component-level spacing; the portal adopts a shell layout with contextual "
        "navigation for contracts and disputes. Screenshots can be pasted in a later revision without altering logo "
        "assets embedded from the original template cover.",
        style="Body Text",
    )

    # --- Implementation ---
    doc.add_page_break()
    doc.add_paragraph("Implementation", style="Heading 2")
    doc.add_paragraph("Technology Stack", style="Heading 3")
    add_table(
        doc,
        ["Layer", "Technologies"],
        [
            ("Client", "Flutter 3.x, Provider, go_router, Supabase Flutter SDK"),
            ("Web", "Next.js App Router, TypeScript, Tailwind patterns as in repo"),
            ("Backend", "Supabase Postgres, Auth, Storage, SQL migrations, RPC"),
            ("Tooling", "GitHub Actions (if enabled), Makefile, Vercel deploy docs"),
        ],
        caption="Table 6.1 — Technology stack",
    )
    doc.add_paragraph("Mobile Application Modules", style="Heading 3")
    doc.add_paragraph(
        "Key modules include `ProposalRepository` (remote sync, deliverable upload, signed URLs), messaging "
        "repositories, jobs provider, onboarding and profile flows, and escrow-related UI on proposal detail and "
        "delivery review routes.",
        style="Body Text",
    )
    doc.add_paragraph("Cloud Functions and Notification Pipeline", style="Heading 3")
    doc.add_paragraph(
        "Server-side logic is expressed as SQL functions where possible, keeping orchestration close to data. "
        "Push notification hooks may be extended; the current MVP emphasises data integrity and portal workflows.",
        style="Body Text",
    )
    doc.add_paragraph("Escrow and Delivery Lifecycle", style="Heading 3")
    add_figure(doc, fig_esc, "Figure 3 — Contract / escrow lifecycle (simplified)")
    doc.add_paragraph("DevOps, Build and Distribution", style="Heading 3")
    doc.add_paragraph(
        "The README documents `make` targets; Flutter web builds integrate with deployment documentation under "
        "`DEPLOY.md`. Environment variables for Supabase keys are excluded from version control per `.gitignore`.",
        style="Body Text",
    )

    # --- Testing ---
    doc.add_page_break()
    doc.add_paragraph("Testing", style="Heading 2")
    doc.add_paragraph("Test Strategy", style="Heading 3")
    doc.add_paragraph(
        "Testing mixes automated unit coverage for repository logic with manual two-device demo passes documented in "
        "`DEMO_NOTES.md`, including timed payout simulation controls.",
        style="Body Text",
    )
    doc.add_paragraph("Unit and Integration Tests", style="Heading 3")
    doc.add_paragraph(
        "Dart tests under `client/test/` validate selected repository behaviours; additional integration tests can "
        "target Supabase with ephemeral databases in CI when secrets are available.",
        style="Body Text",
    )
    doc.add_paragraph("User Acceptance Tests", style="Heading 3")
    doc.add_paragraph(
        "Acceptance scenarios mirror employer and freelancer journeys on web and mobile, verifying that lifecycle "
        "labels, Turkish UI strings where applicable, and download links behave consistently after refresh.",
        style="Body Text",
    )

    # --- Working product ---
    doc.add_page_break()
    doc.add_paragraph("Working Product and MVP", style="Heading 2")
    doc.add_paragraph("Demo Scenario — End-to-End Flow", style="Heading 3")
    doc.add_paragraph(
        "The repository ships `DEMO_NOTES.md` with credential tables and a two-device walkthrough: client accepts a "
        "proposal to fund escrow, freelancer submits delivery, client accepts delivery to enter payout window, and "
        "demo controls advance the clock for presentation purposes.",
        style="Body Text",
    )
    doc.add_paragraph("Working Product Capabilities", style="Heading 3")
    add_table(
        doc,
        ["Capability", "Status (MVP)"],
        [
            ("Job posting and discovery", "Implemented"),
            ("Proposal submit / accept", "Implemented"),
            ("Private deliverable storage + signed download", "Implemented"),
            ("Lifecycle RPCs for delivery review", "Implemented"),
            ("Admin dispute resolution UI", "Implemented (portal)"),
        ],
        caption="Table 8.1 — Working product capability matrix",
    )

    # --- Conclusion ---
    doc.add_page_break()
    doc.add_paragraph("Conclusion and Future Work", style="Heading 2")
    doc.add_paragraph("Overall Evaluation", style="Heading 3")
    doc.add_paragraph(
        "The project demonstrates that a small team can converge a credible marketplace slice by anchoring trust "
        "mechanisms in the database and keeping clients thin. The remaining gap is production hardening: payments "
        "compliance, observability, and richer notification channels.",
        style="Body Text",
    )
    doc.add_paragraph("Lessons Learned", style="Heading 3")
    doc.add_paragraph(
        "Early investment in migrations and RLS reduces rework; keeping demo credentials and repair SQL documented "
        "accelerates onboarding for evaluators.",
        style="Body Text",
    )
    doc.add_paragraph("Future Work", style="Heading 3")
    doc.add_paragraph(
        "Planned extensions include real payment provider integration, automated CI against a disposable Supabase "
        "branch, mobile push parity, and deeper analytics on proposal quality.",
        style="Body Text",
    )

    doc.add_page_break()
    doc.add_paragraph("References", style="Heading 2")
    refs = [
        "Flutter documentation — https://docs.flutter.dev/",
        "Supabase documentation — https://supabase.com/docs",
        "Next.js documentation — https://nextjs.org/docs",
        "OWASP ASVS (overview) — https://owasp.org/www-project-application-security-verification-standard/",
        "PostgreSQL Row Security Policies — https://www.postgresql.org/docs/current/ddl-rowsecurity.html",
    ]
    for r in refs:
        doc.add_paragraph(r, style="Body Text")

    doc.add_page_break()
    doc.add_paragraph("Appendix A — GitHub and Project Management Links", style="Heading 2")
    doc.add_paragraph(
        "Repository: https://github.com/OzgurBuyukikiz01/Prolance_Freelance_App",
        style="Body Text",
    )
    doc.add_paragraph("Appendix B — Installation and Run Instructions", style="Heading 2")
    doc.add_paragraph(
        "See repository `README.md`: `cd client && flutter pub get && flutter run`; `cd web && npm install && npm run dev`; "
        "Supabase project ref documented in README for hosted demo alignment.",
        style="Body Text",
    )
    doc.add_paragraph("Appendix C — Development Tools Used", style="Heading 2")
    doc.add_paragraph(
        "Visual Studio Code / Cursor, Flutter SDK, Node.js LTS, Supabase CLI (optional), Git, GitHub, Chrome device "
        "emulation for Flutter web.",
        style="Body Text",
    )

    doc.save(str(OUT_DOCX))
    print("Wrote", OUT_DOCX)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
